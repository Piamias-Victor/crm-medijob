#!/usr/bin/env python3
"""Génère docs/testing-manual-v1.docx — import Google Docs."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "testing-manual-v1.docx"

# MediJob-ish teal
TEAL = RGBColor(0x0D, 0x73, 0x6B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
OK_GREEN = "C6EFCE"
KO_RED = "FFC7CE"
WAIT_YELLOW = "FFF2CC"
TEAL_BG = "0D736B"
LIGHT_TEAL = "E6F4F2"
GRAY_BG = "F3F4F6"
BLUE_BG = "DEEBF7"
ORANGE_BG = "FCE4D6"


def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, *, bold=False, size=11, color=DARK, italic=False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, bold=True, size=26, color=TEAL)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, size=12, color=MUTED)


def h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, bold=True, size=16, color=TEAL)


def h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, bold=True, size=13, color=DARK)


def para(doc: Document, text: str, *, italic=False, color=DARK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=11, color=color, italic=italic)


def tip(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    shade(cell, LIGHT_TEAL)
    p = cell.paragraphs[0]
    r = p.add_run("💡  " + text)
    set_run(r, size=10, color=TEAL)
    doc.add_paragraph()


def warn(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade(cell, WAIT_YELLOW)
    p = cell.paragraphs[0]
    r = p.add_run("⏳  " + text)
    set_run(r, size=10, color=DARK)
    doc.add_paragraph()


def deferred(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade(cell, ORANGE_BG)
    p = cell.paragraphs[0]
    r = p.add_run("🚧  DIFFÉRÉ — " + text)
    set_run(r, size=10, color=DARK, bold=True)
    doc.add_paragraph()


def check(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run("☐  " + text)
    set_run(r, size=11)


def step(doc: Document, n: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{n}.  ")
    set_run(r, bold=True, size=11, color=TEAL)
    r2 = p.add_run(text)
    set_run(r2, size=11)


def banner_row(doc: Document, cells: list[tuple[str, str]]) -> None:
    """cells = [(text, bg_hex), ...]"""
    table = doc.add_table(rows=1, cols=len(cells))
    table.autofit = True
    for i, (text, bg) in enumerate(cells):
        cell = table.rows[0].cells[i]
        shade(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run(r, bold=True, size=10, color=WHITE if bg == TEAL_BG else DARK)
    doc.add_paragraph()


def simple_table(doc: Document, headers: list[str], rows: list[list[str]], header_bg=TEAL_BG) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, header_bg)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run(r, bold=True, size=10, color=WHITE)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                shade(cell, GRAY_BG)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            set_run(r, size=10)
    doc.add_paragraph()


def role_matrix_cell(cell, value: str) -> None:
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    v = value.strip().lower()
    if v in ("oui", "ok", "✓"):
        shade(cell, OK_GREEN)
        r = p.add_run("OUI")
        set_run(r, bold=True, size=9, color=RGBColor(0x00, 0x61, 0x00))
    elif v in ("non", "✗"):
        shade(cell, KO_RED)
        r = p.add_run("NON")
        set_run(r, bold=True, size=9, color=RGBColor(0x9C, 0x00, 0x06))
    elif "noter" in v or "écart" in v or "attendu" in v:
        shade(cell, WAIT_YELLOW)
        r = p.add_run(value)
        set_run(r, size=8)
    else:
        r = p.add_run(value)
        set_run(r, size=8)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # ——— COVER ———
    add_title(doc, "CRM MediJob — Guide de tests manuels")
    add_subtitle(doc, "Version V1 opérationnelle  ·  Preview Vercel  ·  3 août 2026")
    doc.add_paragraph()

    banner_row(
        doc,
        [
            ("1. Ouvrir preview", TEAL_BG),
            ("2. Se connecter", TEAL_BG),
            ("3. Cocher les ☐", TEAL_BG),
            ("4. Noter les bugs", TEAL_BG),
        ],
    )

    para(
        doc,
        "Ce document est fait pour être rempli pendant les tests. "
        "Coche chaque case ☐ quand c’est OK. Si KO → section Journal de bugs en bas.",
    )
    tip(
        doc,
        "Astuce Google Docs : File → Open → Upload ce fichier .docx, "
        "ou glisse-le dans Drive puis « Ouvrir avec Google Docs ». "
        "Tu peux cocher en remplaçant ☐ par ☑.",
    )

    h2(doc, "Légende")
    simple_table(
        doc,
        ["Symbole", "Signification"],
        [
            ["☐", "À tester"],
            ["☑", "OK — validé"],
            ["✗", "KO — noter dans le journal"],
            ["N/A", "Pas testable sur cet environnement"],
            ["🚧 DIFFÉRÉ", "Feature pas encore livrée (#226, #230, #231, #234)"],
        ],
    )

    # ——— SESSION ———
    h1(doc, "A. Infos de session (à remplir)")
    simple_table(
        doc,
        ["Champ", "Valeur"],
        [
            ["Date", ""],
            ["Testeur", ""],
            ["URL preview Vercel", ""],
            ["Seed DB à jour ?", "Oui / Non"],
            ["IA (extraction)", "Mock / Réel"],
            ["Verdict final", "GO / NO-GO pour démo client"],
        ],
    )

    # ——— ORDER ———
    h1(doc, "B. Ordre recommandé (suivre dans l’ordre)")
    simple_table(
        doc,
        ["#", "Bloc", "Compte recommandé", "Durée approx."],
        [
            ["1", "Socle (login + nav)", "Les 4 rôles", "10 min"],
            ["2", "Accueil", "Recruteur", "5 min"],
            ["3", "Pharmacies", "Recruteur", "20 min"],
            ["4", "Contacts", "Recruteur", "10 min"],
            ["5", "Candidats", "Recruteur", "25 min"],
            ["6", "Missions + Matching", "Recruteur", "25 min"],
            ["7", "Offres", "Recruteur", "10 min"],
            ["8", "Assistant IA", "Recruteur", "10 min"],
            ["9", "Admin + RGPD", "RH-Admin", "15 min"],
            ["10", "Matrice rôles", "Les 4 rôles", "20 min"],
            ["11", "Parcours E2E complet", "Recruteur", "20 min"],
            ["12", "Sections différées", "—", "plus tard"],
        ],
    )

    # ——— ACCOUNTS ———
    h1(doc, "C. Comptes de test")
    para(doc, "Mots de passe par défaut (sauf si la preview a des mots de passe custom).")
    simple_table(
        doc,
        ["Rôle", "Email", "Mot de passe", "À utiliser pour…"],
        [
            [
                "Direction",
                "direction@medijob.fr",
                "direction-medijob-2026",
                "Admin, soft delete, RGPD, finance",
            ],
            [
                "Recruteur",
                "recruteur@medijob.fr",
                "recruteur-medijob-2026",
                "⭐ Parcours métier principal",
            ],
            [
                "Communication",
                "communication@medijob.fr",
                "communication-medijob-2026",
                "Matrice droits (lecture vs écriture)",
            ],
            [
                "RH-Admin",
                "admin@medijob.fr",
                "admin-medijob-2026",
                "Admin, utilisateurs, RGPD",
            ],
        ],
    )
    h2(doc, "Login smoke")
    for t in [
        "Login Direction OK",
        "Login Recruteur OK",
        "Login Communication OK",
        "Login RH-Admin OK",
        "Mauvais mot de passe → refus",
        "Logout → retour page /login",
    ]:
        check(doc, t)

    # ——— PRECHECKS ———
    h1(doc, "D. Avant de commencer")
    h2(doc, "Accès preview")
    for t in [
        "URL preview connue (Vercel → Deployments → branche dev)",
        "Fenêtre privée / navigateur frais (pas de sessions croisées)",
        "Écran ≥ 1280 px (desktop)",
    ]:
        check(doc, t)

    h2(doc, "Données")
    tip(
        doc,
        "Sans seed, les listes sont vides. Demande un reseed Neon avant de continuer.",
    )
    for t in [
        "DB preview seedée récemment (users + pharmacies + candidats + missions)",
        "Upload documents possible (Blob) — sinon marquer N/A sur les docs",
        "Page /login charge sans erreur 500",
        "Logo / couleurs MediJob (teal) visibles",
    ]:
        check(doc, t)

    h2(doc, "Données seed à retrouver")
    simple_table(
        doc,
        ["Type", "Exemples attendus", "☐"],
        [
            ["Pharmacies", "Bellecour (Lyon), Marais (Paris), Vieux-Port (Marseille Prospect), Capitole (Toulouse)", ""],
            ["Contacts", "Sophie Moreau, Thomas Girard, Isabelle Renard…", ""],
            ["Candidats", "Camille Durand, Lucas Martin, Léa Bernard… (≥ 8)", ""],
            ["Missions", "Au moins 1 À pourvoir, 1 Pourvu, 1 Annulée", ""],
            ["Statut UI", "Label « Client » (pas « Actif ») pour les pharmacies actives", ""],
        ],
    )

    # ——— NAV ———
    h1(doc, "E. Navigation & recherche")
    para(doc, "Compte : Recruteur", italic=True, color=MUTED)
    for t in [
        "Sidebar : Accueil, Candidats, Pharmacies, Contacts, Missions, Offres, Assistant IA",
        "Pas d’entrée Admin (Recruteur)",
        "Recherche « Bellecour » → Pharmacie → ouvre la fiche",
        "Recherche « Camille » → Candidat",
        "Recherche d’une mission (ex. Pharmacien adjoint) fonctionne",
    ]:
        check(doc, t)
    para(doc, "Compte : Direction ou RH-Admin", italic=True, color=MUTED)
    for t in [
        "Entrée Admin visible",
        "Sous-menus : Pipeline, Logiciels, Groupements, Métiers, Rôles contact, Utilisateurs, RGPD",
    ]:
        check(doc, t)

    h2(doc, "Auth avancée")
    for t in [
        "Lien mot de passe oublié → page forgot-password",
        "Soumettre un email seed → confirmation (ou message clair si mail non configuré)",
        "Si email reçu : reset MDP puis login OK — sinon N/A",
    ]:
        check(doc, t)

    # ——— HOME ———
    h1(doc, "1. Accueil / Dashboard")
    banner_row(doc, [("Écran : /accueil", LIGHT_TEAL), ("Rôle : Recruteur", BLUE_BG)])
    for t in [
        "KPI visibles (missions à pourvoir, urgentes, candidatures, taux remplissage…)",
        "Chiffres cohérents avec les listes (ordre de grandeur)",
        "Centre d’alertes présent",
        "Clic sur une alerte → écran pertinent",
        "Actions rapides (si présentes) fonctionnent",
    ]:
        check(doc, t)

    # ——— PHARMACIES ———
    h1(doc, "2. Pharmacies")
    banner_row(doc, [("Écran : /pharmacies", LIGHT_TEAL), ("Rôle : Recruteur", BLUE_BG)])

    h2(doc, "2.1 Liste")
    for t in [
        "Tableau (pas seulement des cartes)",
        "Colonnes : nom, ville, code postal, statut, groupement, date d’ajout, référent, œil",
        "Statut affiché Client / Prospect / Inactif",
        "Filtres statut, dpt, groupement, LGO, ville, référent…",
        "Filtrer Client → Bellecour visible ; Prospect → Vieux-Port",
        "Vue rapide (œil) : infos clés + lien vers fiche",
        "Toggle carte : pins visibles + filtres",
    ]:
        check(doc, t)

    h2(doc, "2.2 Création")
    for t in [
        "Recherche SIRENE / SIRET préremplit nom, SIRET, adresse, ville, CP",
        "Référent prérempli mais optionnel (vidable)",
        "Enregistrement → fiche créée",
        "Historique : « Fiche créée » automatique",
    ]:
        check(doc, t)

    h2(doc, "2.3 Import CSV")
    for t in [
        "Upload + mapping colonnes + preview",
        "Doublon SIRET ou nom+ville+CP → écran fusion",
        "Import final OK",
    ]:
        check(doc, t)

    h2(doc, "2.4 Fiche pharmacie")
    for t in [
        "Onglets : Infos, Contacts, Besoins, Historique, Documents",
        "Édition Infos + save → « Fiche modifiée » dans l’historique",
        "Onglet Besoins : missions non terminales + type de contrat visible",
        "Historique : ActivityLog + missions Pourvu / Annulée",
        "Documents : upload PDF + aperçu in-app + téléchargement",
        "Soft delete : Recruteur bloqué ; Direction/RH-Admin OK avec confirmation",
    ]:
        check(doc, t)

    # ——— CONTACTS ———
    h1(doc, "3. Contacts")
    banner_row(doc, [("Écran : /contacts", LIGHT_TEAL), ("Rôle : Recruteur", BLUE_BG)])
    h2(doc, "3.1 Liste")
    for t in [
        "Colonnes nom et prénom séparés",
        "Fonction, pharmacie, tél, email, date, badge principal, œil",
        "Filtres rôle / pharmacie / ville / référent / principal",
        "Vue rapide OK",
    ]:
        check(doc, t)
    h2(doc, "3.2 Création & fiche")
    for t in [
        "Pharmacie obligatoire à la création",
        "Rôle depuis référentiel (ex. Comptabilité si présent)",
        "Référent optionnel",
        "Fiche : édition + historique + documents + soft delete selon rôle",
    ]:
        check(doc, t)

    # ——— CANDIDATES ———
    h1(doc, "4. Candidats (CVthèque)")
    banner_row(doc, [("Écran : /candidats", LIGHT_TEAL), ("Rôle : Recruteur", BLUE_BG)])
    h2(doc, "4.1 Liste")
    for t in [
        "Colonnes : nom, prénom, métier, ville, dpt, référent, dispo, date, statut, œil",
        "Filtres métier / dispo / ville / mobilité / statut / référent…",
        "Vue rapide + toggle carte",
        "Export CSV (si bouton) des lignes filtrées",
        "Zone / onglet candidatures reçues visible",
    ]:
        check(doc, t)
    h2(doc, "4.2 Création / imports")
    for t in [
        "Création manuelle : statut (défaut Nouveau) + prétentions salaire",
        "Doublon email / nom+tél → alerte ou fusion",
        "Import CV (PDF) → revue extraction → création",
        "Import CSV candidats → mapping → doublons gérés",
    ]:
        check(doc, t)
    h2(doc, "4.3 Fiche candidat")
    for t in [
        "Édition profil complète",
        "Bandeau profil incomplet si ville/CP/mobilité/dispo manquants",
        "Positionnement sur mission → statut passe En mission (auto)",
        "Blacklisté reste blacklisté même avec mission",
        "Résumé IA : générer / éditer / sauver",
        "Profil anonymisé + export PDF",
        "Présenter à une pharmacie (mailto / draft)",
        "Historique + onglet missions / pipeline",
        "Documents : upload + aperçu",
        "Effacement RGPD : visible Direction/RH-Admin seulement",
    ]:
        check(doc, t)
    deferred(doc, "Fiche entretien HITL (#226) — à retester plus tard")

    # ——— MISSIONS ———
    h1(doc, "5. Missions & Matching")
    banner_row(doc, [("Écran : /missions", LIGHT_TEAL), ("Rôle : Recruteur", BLUE_BG)])
    h2(doc, "5.1 Liste")
    for t in [
        "Tableau : intitulé, métier, pharmacie, ville, statut, référent, date, œil",
        "Toggle kanban optionnel",
        "Filtres contrat / statut / métier / ville / pharmacie / période",
        "Vue rapide + carte (pins pharmacie)",
    ]:
        check(doc, t)
    h2(doc, "5.2 Création & fiche")
    for t in [
        "Création avec titre, description, profil recherché, contrat, pharmacie, dates…",
        "Édition fiche + pipeline candidats",
        "Changer statut (À pourvoir → Pourvu / Annulée)",
        "Documents + historique + soft delete selon rôle",
    ]:
        check(doc, t)
    h2(doc, "5.3 Matching")
    for t in [
        "Lancer suggestions → liste scorée",
        "Candidat avec prétentions apparaît / score cohérent",
        "Multi-sélection 2 candidats",
        "Actions email / SMS / WhatsApp (deep links)",
        "Ajouter au pipeline → stage initial",
    ]:
        check(doc, t)

    # ——— OFFERS ———
    h1(doc, "6. Offres")
    banner_row(doc, [("Écran : /offres", LIGHT_TEAL), ("Rôle : Recruteur", BLUE_BG)])
    for t in [
        "Vraie liste (plus de placeholder)",
        "Colonnes : titre, statut, mission, nb candidatures",
        "Depuis une mission : générer offre IA → brouillon",
        "Publier → statut PUBLIEE ; Dépublier → DEPUBLIEE",
    ]:
        check(doc, t)
    deferred(doc, "Sync Webflow / CMS (#230) + candidatures site (#231)")

    # ——— ASSISTANT ———
    h1(doc, "7. Assistant IA")
    banner_row(doc, [("Écran : /assistant", LIGHT_TEAL), ("Rôle : Recruteur", BLUE_BG)])
    for t in [
        "Chat libre répond (mock OK)",
        "Raccourcis résumé candidat / pharmacie",
        "Raccourcis mail candidat / pharmacie",
        "Générer offre (contexte mission)",
        "Rapport semaine : chiffres cohérents (pas texte vide)",
        "Meilleurs profils (contexte mission) s’appuie sur matching",
    ]:
        check(doc, t)

    # ——— ADMIN ———
    h1(doc, "8. Admin & RGPD")
    banner_row(doc, [("Écran : /admin", LIGHT_TEAL), ("Rôle : RH-Admin / Direction", BLUE_BG)])
    for t in [
        "Recruteur / Communication : Admin interdit ou redirect",
        "CRUD Pipeline (ordre inclus)",
        "CRUD Logiciels LGO",
        "CRUD Groupements",
        "CRUD Métiers + matrice compatibilité",
        "CRUD Rôles contact (créer un rôle test, l’utiliser)",
        "Utilisateurs : créer, changer rôle, soft delete (garder ≥ 1 admin)",
        "Page RGPD accessible",
        "Lien registre externe si configuré",
        "Trace des effacements visible si UI présente",
    ]:
        check(doc, t)

    # ——— ROLE MATRIX ———
    h1(doc, "9. Matrice des rôles (obligatoire)")
    para(
        doc,
        "Connecte-toi avec chaque rôle et vérifie la case. "
        "Vert = autorisé · Rouge = interdit · Jaune = à noter (écart possible vs spec).",
    )

    headers = ["Scénario", "Direction", "RH-Admin", "Recruteur", "Communication"]
    rows_data = [
        ["Voir menu Admin", "oui", "oui", "non", "non"],
        ["Créer / éditer CRM", "oui", "oui", "oui", "Noter écart (souvent write)"],
        ["Soft delete entité", "oui", "oui", "non", "non"],
        ["Export CSV (si gated)", "oui", "oui", "non", "non"],
        ["Voir CA / Marge", "oui si UI", "oui si UI", "non", "non"],
        ["Effacement RGPD", "oui", "oui", "non", "non"],
        ["Publier une offre", "oui", "oui", "oui", "oui"],
        ["Matching + rattacher", "oui", "oui", "oui", "Attendu NON"],
    ]
    table = doc.add_table(rows=1 + len(rows_data), cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, TEAL_BG)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run(r, bold=True, size=9, color=WHITE)
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ci == 0:
                if ri % 2 == 1:
                    shade(cell, GRAY_BG)
                p = cell.paragraphs[0]
                r = p.add_run(val)
                set_run(r, size=9, bold=True)
            else:
                role_matrix_cell(cell, val)
    doc.add_paragraph()
    for t in [
        "Matrice Admin OK pour les 4 rôles",
        "Soft delete vérifié",
        "RGPD vérifié",
        "Écarts Communication notés dans le journal",
    ]:
        check(doc, t)

    # ——— E2E ———
    h1(doc, "10. Parcours bout-en-bout (E2E)")
    para(doc, "Compte : Recruteur — enchaîne les étapes sans sauter.", italic=True, color=MUTED)

    h2(doc, "Parcours A — Staffing classique")
    for i, t in enumerate(
        [
            "Ouvrir / créer une pharmacie Client (ex. Bellecour)",
            "Vérifier / ajouter un contact titulaire",
            "Créer une mission CDI avec profil recherché",
            "Lancer matching → choisir un candidat (ex. Camille Durand)",
            "Rattacher au pipeline (+ mailto optionnel)",
            "Avancer d’un stage dans le pipeline",
            "Générer offre IA → publier dans le CRM",
            "Passer la mission en Pourvu",
            "Pharmacie : besoin disparu + historique pourvu",
            "Candidat : statut En mission + historique à jour",
        ],
        start=1,
    ):
        step(doc, i, t)
        check(doc, f"Étape {i} OK")

    h2(doc, "Parcours B — CVthèque")
    for i, t in enumerate(
        [
            "Importer un CV PDF → revue → création",
            "Compléter mobilité / ville si bandeau incomplet",
            "Générer résumé IA + PDF anonymisé",
            "Présenter à une pharmacie (draft mail)",
        ],
        start=1,
    ):
        step(doc, i, t)
        check(doc, f"Étape B{i} OK")

    h2(doc, "Parcours C — Qualité données")
    for i, t in enumerate(
        [
            "Créer candidat avec email déjà existant → doublon géré",
            "Import CSV pharmacie avec SIRET doublon → fusion",
        ],
        start=1,
    ):
        step(doc, i, t)
        check(doc, f"Étape C{i} OK")

    # ——— DEFERRED ———
    h1(doc, "11. À retester plus tard (pas encore livré)")
    deferred(doc, "#226 Fiche entretien HITL")
    check(doc, "Section entretien sur création/fiche")
    check(doc, "Champs sauvegardés sur le profil")

    deferred(doc, "#230 Sync offres → Webflow / CMS")
    check(doc, "Publish crée/maj item site")
    check(doc, "Unpublish retire côté site")

    deferred(doc, "#231 Candidatures site (dépend de #230)")
    check(doc, "Webhook → Application en inbox")
    check(doc, "Accept → Candidate + CV ; Refus → REFUSEE")

    deferred(doc, "#234 Finance (après Q13 client)")
    check(doc, "Écrans perf / facturation / devis")
    check(doc, "CA/Marge visibles Direction & RH-Admin seulement")

    # ——— BUGS ———
    h1(doc, "12. Journal de bugs")
    para(
        doc,
        "Sévérité : S1 = bloque le métier · S2 = contournable · S3 = cosmétique",
        italic=True,
        color=MUTED,
    )
    simple_table(
        doc,
        ["ID", "Rôle", "URL", "Étapes", "Attendu", "Obtenu", "Sév.", "Capture ?"],
        [
            ["B1", "", "", "", "", "", "", ""],
            ["B2", "", "", "", "", "", "", ""],
            ["B3", "", "", "", "", "", "", ""],
            ["B4", "", "", "", "", "", "", ""],
            ["B5", "", "", "", "", "", "", ""],
        ],
    )

    # ——— SUMMARY ———
    h1(doc, "13. Synthèse finale")
    simple_table(
        doc,
        ["Question", "Réponse"],
        [
            ["E2E Staffing (A) OK ?", ""],
            ["Matrice rôles OK ?", ""],
            ["Nb bugs S1 / S2 / S3", ""],
            ["Go démo client ?", "GO / NO-GO"],
            ["Commentaire libre", ""],
        ],
    )

    para(
        doc,
        "Références repo : docs/PRD_V1_OPERATIONNEL.md · docs/grill/QUESTIONS_CLIENT_V1.md · issues #226 #230 #231 #234",
        italic=True,
        color=MUTED,
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
